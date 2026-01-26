"""
Resume Parser: Extract structured data from PDF and DOCX resume files.
Returns JSON with skills, education, experience, projects, and completeness score.
"""

import json
import re
import os
from pathlib import Path
from typing import Dict, List, Tuple, Optional
from datetime import datetime

import sys

# Try to import PDF libraries
try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("Warning: pdfplumber not installed. PDF parsing disabled.", file=sys.stderr)

# Try to import DOCX library
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. DOCX parsing disabled.", file=sys.stderr)


class ResumeParser:
    """Parse resume files and extract structured data."""

    # Comprehensive keyword lists for skill extraction
    PROGRAMMING_LANGUAGES = [
        'python', 'java', 'javascript', 'typescript', 'c#', 'c++', 'c',
        'php', 'ruby', 'go', 'rust', 'kotlin', 'swift', 'objectivec',
        'scala', 'r', 'matlab', 'vb', 'perl', 'groovy', 'dart',
        'elixir', 'clojure', 'haskell', 'lua', 'shell', 'bash'
    ]

    WEB_FRAMEWORKS = [
        'react', 'angular', 'vue', 'svelte', 'next.js', 'nuxt',
        'express', 'django', 'flask', 'fastapi', 'spring', 'spring boot',
        'rails', 'laravel', 'asp.net', 'nestjs', 'fastify'
    ]

    DATABASES = [
        'mysql', 'postgresql', 'mongodb', 'redis', 'elasticsearch',
        'cassandra', 'dynamodb', 'oracle', 'sql server', 'sqlite',
        'firestore', 'couchdb', 'mariadb', 'neo4j', 'influxdb'
    ]

    TOOLS_PLATFORMS = [
        'git', 'github', 'gitlab', 'bitbucket', 'docker', 'kubernetes',
        'jenkins', 'circleci', 'travis', 'gitlab ci', 'aws', 'azure',
        'gcp', 'heroku', 'netlify', 'vercel', 'terraform', 'ansible',
        'prometheus', 'grafana', 'datadog', 'jira', 'confluence'
    ]

    DATA_SCIENCE = [
        'tensorflow', 'pytorch', 'scikit-learn', 'pandas', 'numpy',
        'matplotlib', 'seaborn', 'jupyter', 'spark', 'hadoop',
        'kafka', 'airflow', 'dbt', 'looker', 'tableau', 'power bi'
    ]

    # Degree keywords
    DEGREE_KEYWORDS = {
        'bachelor': ['bachelor', 'b.s.', 'b.a.', 'bs', 'ba', 'undergraduate'],
        'master': ['master', 'm.s.', 'm.a.', 'ms', 'ma', 'mba', 'graduate'],
        'phd': ['phd', 'ph.d.', 'doctorate', 'doctor of philosophy'],
        'diploma': ['diploma', 'associate', 'a.s.', 'as'],
        'bootcamp': ['bootcamp', 'certification', 'certificate']
    }

    # Common education institutions for better extraction
    INSTITUTION_KEYWORDS = [
        'university', 'college', 'institute', 'school', 'academy',
        'polytechnic', 'technical', 'engineering'
    ]

    # Job role keywords to identify experience
    JOB_ROLE_KEYWORDS = [
        'engineer', 'developer', 'programmer', 'designer', 'analyst',
        'manager', 'director', 'lead', 'senior', 'junior', 'intern',
        'consultant', 'specialist', 'architect', 'scientist', 'administrator'
    ]

    # Generic words to filter out from skills (not actual skills)
    SKILL_BLOCKLIST = [
        'technical', 'knowledge', 'skills', 'skill', 'proficient', 'experience',
        'experienced', 'expertise', 'proficiency', 'ability', 'abilities',
        'strong', 'excellent', 'good', 'great', 'advanced', 'intermediate',
        'beginner', 'basic', 'familiar', 'understanding', 'working', 'hands',
        'on', 'and', 'or', 'the', 'with', 'in', 'of', 'for', 'to', 'a', 'an',
        'using', 'used', 'use', 'including', 'such', 'as', 'like', 'etc',
        'various', 'multiple', 'several', 'many', 'other', 'others', 'more',
        'technologies', 'technology', 'tools', 'tool', 'frameworks', 'framework',
        'languages', 'language', 'platforms', 'platform', 'systems', 'system',
        'software', 'hardware', 'applications', 'application', 'solutions',
        'development', 'programming', 'coding', 'building', 'creating',
        'developing', 'designing', 'implementing', 'managing', 'leading',
        'team', 'teams', 'project', 'projects', 'work', 'worked', 'working',
        'company', 'companies', 'organization', 'organizations', 'client',
        'clients', 'customer', 'customers', 'user', 'users', 'data', 'information',
        'problem', 'problems', 'solving', 'solution', 'analysis', 'design',
        'implementation', 'testing', 'deployment', 'maintenance', 'support',
        'communication', 'collaboration', 'leadership', 'management', 'time',
        'learning', 'quick', 'fast', 'efficient', 'effective', 'responsible',
        'responsibility', 'responsibilities', 'duties', 'duty', 'role', 'roles'
    ]

    # Words that indicate something is NOT an institution (false positives)
    INSTITUTION_BLOCKLIST = [
        'hackathon', 'competition', 'contest', 'event', 'conference', 'summit',
        'workshop', 'seminar', 'webinar', 'meetup', 'bootcamp', 'camp', 'fest',
        'festival', 'olympiad', 'challenge', 'award', 'prize', 'winner', 'won',
        'participated', 'participant', 'organized', 'organizer', 'hosted',
        'host', 'speaker', 'presented', 'presentation', 'talk', 'session',
        'track', 'category', 'level', 'round', 'stage', 'phase', 'final',
        'semifinal', 'quarterfinal', 'national', 'international', 'regional',
        'state', 'district', 'city', 'local', 'online', 'virtual', 'remote'
    ]

    def __init__(self, file_path: str):
        """Initialize parser with file path."""
        self.file_path = Path(file_path)
        self.text = ""
        self.lines = []

    def extract_text(self) -> str:
        """Extract text from PDF or DOCX file."""
        if not self.file_path.exists():
            raise FileNotFoundError(f"File not found: {self.file_path}")

        if self.file_path.suffix.lower() == '.pdf':
            if not PDF_AVAILABLE:
                raise ImportError("pdfplumber is required for PDF parsing. Install with: pip install pdfplumber")
            self.text = self._extract_pdf()
        elif self.file_path.suffix.lower() == '.docx':
            if not DOCX_AVAILABLE:
                raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")
            self.text = self._extract_docx()
        else:
            raise ValueError(f"Unsupported file type: {self.file_path.suffix}")

        # Split into lines for easier processing
        self.lines = [line.strip() for line in self.text.split('\n') if line.strip()]
        return self.text

    def _extract_pdf(self) -> str:
        """Extract text from PDF file."""
        text = ""
        with pdfplumber.open(self.file_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
                text += "\n"
        return text

    def _extract_docx(self) -> str:
        """Extract text from DOCX file."""
        doc = Document(self.file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        return text

    def extract_skills(self) -> List[str]:
        """Extract technical skills from resume text."""
        if not self.text:
            self.extract_text()

        skills_found = set()
        text_lower = self.text.lower()

        # Extract programming languages
        for lang in self.PROGRAMMING_LANGUAGES:
            if self._keyword_in_text(lang, text_lower):
                skills_found.add(lang.title())

        # Extract web frameworks
        for framework in self.WEB_FRAMEWORKS:
            if self._keyword_in_text(framework, text_lower):
                skills_found.add(framework.title())

        # Extract databases
        for db in self.DATABASES:
            if self._keyword_in_text(db, text_lower):
                skills_found.add(db.title())

        # Extract tools and platforms
        for tool in self.TOOLS_PLATFORMS:
            if self._keyword_in_text(tool, text_lower):
                skills_found.add(tool.title())

        # Extract data science tools
        for tool in self.DATA_SCIENCE:
            if self._keyword_in_text(tool, text_lower):
                skills_found.add(tool.title())

        # Extract generic skill patterns (noun + skill-related words)
        skill_patterns = [
            r'\b(?:proficient|experienced|skilled|expertise)\s+(?:in|with)?\s+([a-zA-Z\s&,\.]+?)(?:,|and|\.|;|\n)',
            r'\b(?:technical\s+)?skills?:?\s*\n?([a-zA-Z\s&,\.]+?)(?:\n|•)',
        ]

        for pattern in skill_patterns:
            matches = re.finditer(pattern, self.text, re.IGNORECASE | re.MULTILINE)
            for match in matches:
                skills_text = match.group(1)
                skill_items = [s.strip() for s in re.split(r'[,&•]', skills_text)]
                for skill in skill_items:
                    if len(skill) > 2 and len(skill) < 50:
                        skills_found.add(skill.strip())

        # Filter out generic/non-skill words
        filtered_skills = []
        for skill in skills_found:
            skill_lower = skill.lower().strip()
            # Skip if it's in the blocklist
            if skill_lower in self.SKILL_BLOCKLIST:
                continue
            # Skip if it's too short (single char or two chars)
            if len(skill_lower) < 2:
                continue
            # Skip if it contains only numbers
            if skill_lower.isdigit():
                continue
            # Skip if it's mostly whitespace or punctuation
            if len(re.sub(r'[\s\W]', '', skill_lower)) < 2:
                continue
            filtered_skills.append(skill)

        return sorted(list(set(filtered_skills)))

    def extract_education(self) -> List[Dict[str, Optional[str]]]:
        """Extract education information from resume."""
        if not self.text:
            self.extract_text()

        education_list = []
        text_lower = self.text.lower()

        # Find degree types
        for degree_type, keywords in self.DEGREE_KEYWORDS.items():
            for keyword in keywords:
                # Find lines containing degree keyword
                for i, line in enumerate(self.lines):
                    if keyword.lower() in line.lower():
                        degree_info = {
                            'degree': degree_type.title(),
                            'institution': None,
                            'year': None
                        }

                        # Look for institution in this line and nearby lines
                        context = ' '.join(self.lines[max(0, i-2):min(len(self.lines), i+3)])
                        institution = self._extract_institution(context)
                        if institution:
                            degree_info['institution'] = institution

                        # Extract year and normalize to 4-digit format
                        year = self._extract_year(line)
                        degree_info['year'] = self._normalize_year(year)

                        # Avoid duplicates
                        if degree_info not in education_list:
                            education_list.append(degree_info)

        # Clean up education entries
        cleaned_education = self._clean_education_entries(education_list)
        return cleaned_education

    def extract_experience_months(self) -> int:
        """Calculate total months of experience from resume."""
        if not self.text:
            self.extract_text()

        total_months = 0
        experience_entries = []

        # Pattern to find date ranges (e.g., "Jan 2020 - Dec 2021")
        date_patterns = [
            r'(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+(\d{4})\s*[-–]\s*(present|jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+(\d{4})?',
            r'(\d{1,2})/(\d{1,2})/(\d{4})\s*[-–]\s*(present|\d{1,2}/\d{1,2}/\d{4})',
            r'(\d{4})\s*[-–]\s*(present|\d{4})',
        ]

        for pattern in date_patterns:
            matches = re.finditer(pattern, self.text, re.IGNORECASE)
            for match in matches:
                try:
                    start_month, start_year, end_part, end_year = self._parse_date_range(match)
                    
                    # Calculate months
                    if end_year:
                        months = (int(end_year) - int(start_year)) * 12
                        months += end_month - start_month
                        experience_entries.append(months)
                    elif 'present' in end_part.lower():
                        # Assume work until now
                        current_year = datetime.now().year
                        months = (current_year - int(start_year)) * 12
                        experience_entries.append(months)
                except (ValueError, IndexError):
                    continue

        # Sum up all experiences (removing outliers and duplicates)
        if experience_entries:
            # Filter out unrealistic entries (> 50 years or < 1 month)
            valid_entries = [m for m in experience_entries if 1 <= m <= 600]
            total_months = max(valid_entries) if valid_entries else sum(valid_entries) // len(valid_entries) if valid_entries else 0

        return total_months

    def extract_projects_count(self) -> int:
        """Count number of projects mentioned in resume."""
        if not self.text:
            self.extract_text()

        projects_count = 0

        # Look for "Projects" or "Portfolio" sections
        section_pattern = r'(?:projects?|portfolio|portfolio\s+projects?)\s*:?\s*\n'
        sections = re.split(section_pattern, self.text, flags=re.IGNORECASE)

        if len(sections) > 1:
            projects_section = sections[1]

            # Count bullet points or numbered items
            bullet_pattern = r'(?:^|\n)\s*[-•*]\s+'
            bullets = re.findall(bullet_pattern, projects_section)
            projects_count += len(bullets)

            # Count numbered items
            numbered_pattern = r'(?:^|\n)\s*\d+\.\s+'
            numbered = re.findall(numbered_pattern, projects_section)
            projects_count += len(numbered)

        # Also look for project-related keywords
        project_keywords = ['built', 'developed', 'created', 'designed', 'implemented']
        for keyword in project_keywords:
            # Count occurrences in experience section
            pattern = rf'\b{keyword}\b'
            matches = re.findall(pattern, self.text, re.IGNORECASE)
            # Only count if it's in a work context (near job titles)
            if len(matches) > 0:
                projects_count += min(len(matches) // 2, 1)  # Avoid over-counting

        return min(projects_count, 20)  # Cap at 20 to avoid unrealistic numbers

    def calculate_completeness_score(self) -> float:
        """Calculate resume completeness score (0-1)."""
        if not self.text:
            self.extract_text()

        score = 0.0
        max_score = 10.0

        # Check for contact information
        if re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', self.text):
            score += 1.5  # Email found

        if re.search(r'\b(?:\+?1)?[\s.-]?\(?[2-9]\d{2}\)?[\s.-]?[2-9]\d{2}[\s.-]?\d{4}\b', self.text):
            score += 0.5  # Phone number found

        # Check for sections
        sections = ['experience', 'education', 'skills', 'projects', 'portfolio']
        sections_found = 0
        for section in sections:
            if re.search(rf'\b{section}\b', self.text, re.IGNORECASE):
                sections_found += 1

        score += (sections_found / len(sections)) * 3  # Up to 3 points for sections

        # Check for education
        if self.extract_education():
            score += 1.5

        # Check for skills
        if self.extract_skills():
            score += 1.5

        # Check for experience
        experience_months = self.extract_experience_months()
        if experience_months > 0:
            score += 1.5

        # Check for projects
        projects = self.extract_projects_count()
        if projects > 0:
            score += 0.5

        # Check resume length (should be meaningful)
        word_count = len(self.text.split())
        if word_count > 200:
            score += 1  # Reasonable length

        # Normalize to 0-1 scale
        completeness = min(score / max_score, 1.0)
        return round(completeness, 2)

    def _keyword_in_text(self, keyword: str, text: str) -> bool:
        """Check if keyword appears in text as whole word."""
        pattern = rf'\b{re.escape(keyword)}\b'
        return bool(re.search(pattern, text, re.IGNORECASE))

    def _extract_institution(self, text: str) -> Optional[str]:
        """Extract institution name from text."""
        # Look for university/college names
        patterns = [
            r'(?:at|from)?\s+([A-Z][a-zA-Z\s&\-\.]+(?:University|College|Institute|School|Academy))',
            r'([A-Z][a-zA-Z\s&\-\.]+(?:University|College|Institute|School|Academy))',
        ]

        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                institution = match.group(1).strip()
                if len(institution) > 3 and len(institution) < 100:
                    return institution

        return None

    def _extract_year(self, text: str) -> Optional[str]:
        """Extract graduation year from text."""
        year_pattern = r'\b(19|20)\d{2}\b'
        matches = re.findall(year_pattern, text)
        if matches:
            # Return the last year found (likely graduation year)
            return matches[-1]
        return None

    def _normalize_year(self, year: Optional[str]) -> Optional[int]:
        """Normalize year to 4-digit integer format."""
        if year is None:
            return None
        
        # If it's already a 4-digit string matching 19xx or 20xx
        if isinstance(year, str):
            # Try to find a 4-digit year
            match = re.search(r'(19|20)\d{2}', year)
            if match:
                return int(match.group(0))
            # Try 2-digit year (assume 20xx for < 50, 19xx for >= 50)
            match = re.search(r'\b(\d{2})\b', year)
            if match:
                two_digit = int(match.group(1))
                if two_digit < 50:
                    return 2000 + two_digit
                else:
                    return 1900 + two_digit
        elif isinstance(year, int):
            if 1900 <= year <= 2100:
                return year
        
        return None

    def _clean_education_entries(self, education_list: List[Dict]) -> List[Dict]:
        """Clean and deduplicate education entries."""
        seen = set()
        cleaned = []
        
        for entry in education_list:
            degree = entry.get('degree', '')
            institution = entry.get('institution')
            year = entry.get('year')
            
            # Skip entries with no institution
            if institution is None or institution.strip() == '':
                continue
            
            # Skip if institution looks like a non-institution (event, competition, etc.)
            inst_lower = institution.lower()
            is_blocklisted = False
            for blocked in self.INSTITUTION_BLOCKLIST:
                if blocked in inst_lower:
                    is_blocklisted = True
                    break
            if is_blocklisted:
                continue
            
            # Skip if institution is too long (likely a sentence, not an institution name)
            if len(institution) > 80:
                continue
            
            # Skip if institution doesn't contain any institution keyword
            has_institution_keyword = False
            for keyword in self.INSTITUTION_KEYWORDS:
                if keyword.lower() in inst_lower:
                    has_institution_keyword = True
                    break
            if not has_institution_keyword:
                continue
            
            # Create deduplication key (degree + institution, case-insensitive)
            dedup_key = (degree.lower().strip(), inst_lower.strip())
            if dedup_key in seen:
                continue
            seen.add(dedup_key)
            
            # Add cleaned entry
            cleaned.append({
                'degree': degree,
                'institution': institution.strip(),
                'year': year
            })
        
        return cleaned

    def _parse_date_range(self, match) -> Tuple[int, int, str, Optional[int]]:
        """Parse date range from regex match."""
        groups = match.groups()

        # Handle different date formats
        if len(groups) >= 4:
            start_part = groups[0]
            start_year = int(groups[1])
            end_part = groups[2]
            end_year = int(groups[3]) if groups[3] else None

            # Convert month string to number
            months = {
                'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
                'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12
            }

            start_month = 1
            end_month = 12

            if isinstance(start_part, str):
                month_match = re.match(r'([a-z]+)', start_part, re.IGNORECASE)
                if month_match:
                    start_month = months.get(month_match.group(1).lower()[:3], 1)

            if isinstance(end_part, str) and 'present' not in end_part.lower():
                month_match = re.match(r'([a-z]+)', end_part, re.IGNORECASE)
                if month_match:
                    end_month = months.get(month_match.group(1).lower()[:3], 12)

            return start_month, start_year, end_part, end_year

        return 1, 0, "", None

    def parse(self) -> Dict:
        """Parse resume and return structured data as JSON."""
        self.extract_text()

        # DEBUG: Log extracted text information (stderr only)
        text_length = len(self.text)
        text_preview = self.text[:300] if self.text else ""
        
        print(f"\n[DEBUG] Resume Text Extraction Complete", file=sys.stderr)
        print(f"[DEBUG] Extracted text length: {text_length} characters", file=sys.stderr)
        print(f"[DEBUG] First 300 characters: {text_preview}", file=sys.stderr)
        
        if text_length == 0:
            print("[WARNING] ⚠️  EMPTY RESUME - No text was extracted from the file!", file=sys.stderr)
        elif text_length < 100:
            print("[WARNING] ⚠️  VERY SHORT RESUME - Only {0} characters extracted. File may be empty or unreadable.".format(text_length), file=sys.stderr)
        
        result = {
            'skills': self.extract_skills(),
            'education': self.extract_education(),
            'experience_months': self.extract_experience_months(),
            'projects_count': self.extract_projects_count(),
            'resume_completeness_score': self.calculate_completeness_score()
        }

        return result

    def to_json(self) -> str:
        """Return parsed data as JSON string."""
        return json.dumps(self.parse(), indent=2)


def parse_resume(file_path: str) -> Dict:
    """Convenience function to parse resume and return data."""
    parser = ResumeParser(file_path)
    return parser.parse()


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python resume_parser.py <resume_file.pdf|resume_file.docx>", file=sys.stderr)
        print("\nExample: python resume_parser.py resume.pdf", file=sys.stderr)
        sys.exit(1)

    file_path = sys.argv[1]

    try:
        result = parse_resume(file_path)
        # ONLY valid JSON goes to stdout
        print(json.dumps(result))
    except FileNotFoundError as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)
    except ImportError as e:
        print(f"Error: {e}", file=sys.stderr)
        print("\nInstall required packages with:", file=sys.stderr)
        print("  pip install pdfplumber python-docx", file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print(f"Error parsing resume: {e}", file=sys.stderr)
        sys.exit(1)
